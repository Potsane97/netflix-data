import streamlit as st
from PIL import Image, ImageOps, ImageFilter
import pytesseract
from io import BytesIO
from docx import Document

# Ensure pytesseract knows where tesseract binary is (default on Debian installs)
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

st.set_page_config(page_title="Image → Word", layout="centered")
st.title("🖼️ → 📝 Image to Word")

uploaded_files = st.file_uploader("Upload image(s)", type=["png","jpg","jpeg","tiff"], accept_multiple_files=True)

def preprocess_pil(img: Image.Image) -> Image.Image:
    # convert to grayscale, increase contrast, and optionally sharpen
    img = img.convert("L")
    img = ImageOps.autocontrast(img)
    img = img.filter(ImageFilter.SHARPEN)
    return img

def image_to_text(img: Image.Image) -> str:
    pil = preprocess_pil(img)
    # choose a reasonable PSM for documents//
    custom_config = r'--oem 3 --psm 3'
    text = pytesseract.image_to_string(pil, lang='eng', config=custom_config)
    return text

if uploaded_files:
    st.write(f"Found {len(uploaded_files)} file(s). Processing...")
    doc = Document()
    for idx, uploaded_file in enumerate(uploaded_files, start=1):
        # load image
        img = Image.open(uploaded_file).convert("RGB")
        st.image(img, caption=f"Input {idx}", use_column_width=True)
        text = image_to_text(img)
        if text.strip() == "":
            st.warning(f"No text found in file {uploaded_file.name}.")
        else:
            st.subheader(f"Extracted text — {uploaded_file.name}")
            st.text_area("Preview", value=text, height=200)
            # Add heading and text to doc
            doc.add_heading(uploaded_file.name, level=2)
            for para in text.splitlines():
                if para.strip():
                    doc.add_paragraph(para)
            doc.add_page_break()

    # Save doc to bytes and provide download
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("📄 Download Word (.docx)", data=buffer, file_name="output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Upload one or more images to extract text and generate a Word document.")
